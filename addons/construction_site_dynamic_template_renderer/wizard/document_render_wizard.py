from odoo import models, fields, api
from odoo.exceptions import ValidationError
from docx import Document
from io import BytesIO
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docxcompose.composer import Composer
from docx.enum.section import WD_SECTION
from tempfile import NamedTemporaryFile
from docx2pdf import convert
from PyPDF2 import PdfReader
import base64
import zipfile
import os
import subprocess


class DocumentRenderWizard(models.TransientModel):
    _name = 'document.render.wizard'
    _description = 'Document Render Wizard'

    project_id = fields.Many2one('project.project', string="Project", required=True)
    task_id = fields.Many2one('project.task', string="Task")
    document_template_id = fields.Many2one('document.template', string="Document Template")
    file_selection_ids = fields.One2many(
        'document.render.wizard.file.selection',
        'wizard_id',
        string="Available Files"
    )

    def action_select_all(self):
        self.file_selection_ids.write({'selected': True})
        return {
            'type': 'ir.actions.act_window',
            'res_model': self._name,
            'view_mode': 'form',
            'res_id': self.id,
            'target': 'new',
            'view_id': self.env.ref('construction_site_dynamic_template_renderer.view_document_render_wizard_form').id,
        }

    def action_deselect_all(self):
        self.file_selection_ids.write({'selected': False})
        return {
            'type': 'ir.actions.act_window',
            'res_model': self._name,
            'view_mode': 'form',
            'res_id': self.id,
            'target': 'new',
            'view_id': self.env.ref('construction_site_dynamic_template_renderer.view_document_render_wizard_form').id,
        }

    @api.model
    def default_get(self, fields_list):
        res = super(DocumentRenderWizard, self).default_get(fields_list)
        active_id = self.env.context.get('active_id')
        active_model = self.env.context.get('active_model')

        project_id = False
        record_task_id = False

        if active_id and active_model:
            if active_model == 'project.task':
                task = self.env['project.task'].browse(active_id)
                project_id = task.project_id.id if task.project_id else False
                record_task_id = task.id
            elif active_model == 'project.project':
                project_id = active_id

        file_selections = []
        if project_id:
            project_id = self.env['project.project'].browse(active_id)
            for task_id in project_id.task_ids:
                for opis_param in task_id.opis_parameter_ids:
                    if opis_param.document_sablon_id:
                        file_selections.append((0, 0, {
                            'name': opis_param.name,
                            'denumire_document_opis': opis_param.denumire_document_opis,
                            'prefix_document': opis_param.prefix_document,
                            'cod_document': opis_param.cod_document,
                            'data_document': opis_param.data_document,
                            'attachment_id': opis_param.document_sablon_id.id,
                            'project_id': project_id.id,
                            'task_id': task_id.id,
                            'document_template_id': opis_param.document_sablon_id.res_id,
                        }))

        res.update({
            'project_id': project_id.id,
            'file_selection_ids': file_selections,
        })

        if record_task_id:
            res.update({'task_id': record_task_id.id})

        return res

    def replace_placeholder_across_runs(self, paragraph, replacements):
        runs = paragraph.runs
        if not runs:
            return

        full_text = "".join(run.text for run in runs)

        for key, value in replacements.items():
            start_idx = full_text.find(key)
            while start_idx != -1:
                end_idx = start_idx + len(key)

                char_count = 0
                run_indexes = []
                for i, run in enumerate(runs):
                    run_len = len(run.text)
                    run_start = char_count
                    run_end = char_count + run_len
                    char_count += run_len

                    if run_end > start_idx and run_start < end_idx:
                        run_indexes.append(i)

                if not run_indexes:
                    break

                first_run = runs[run_indexes[0]]
                last_run = runs[run_indexes[-1]]

                involved_text = "".join(runs[i].text for i in run_indexes)
                replaced_text = involved_text.replace(key, value)
                first_run.text = replaced_text

                for i in run_indexes[1:]:
                    runs[i].text = ""

                full_text = full_text[:start_idx] + value + full_text[end_idx:]
                start_idx = full_text.find(key, start_idx + len(value))

    def add_page_number_footer(self, doc):
        section = doc.sections[0]
        footer = section.footer
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        paragraph.alignment = 1

        run = paragraph.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')

        instrText = OxmlElement('w:instrText')
        instrText.text = 'PAGE'

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

    def get_pdf_page_count_from_docx(self, docx_bytes):
        with NamedTemporaryFile(suffix=".docx", delete=False) as tmp_docx:
            tmp_docx.write(docx_bytes)
            tmp_docx.flush()
            tmp_docx_path = tmp_docx.name

        pdf_path = tmp_docx_path.replace(".docx", ".pdf")
        subprocess.run([
            "libreoffice", "--headless", "--convert-to", "pdf",
            "--outdir", os.path.dirname(tmp_docx_path), tmp_docx_path
        ], check=True)

        with open(pdf_path, "rb") as f_pdf:
            pdf_reader = PdfReader(f_pdf)
            num_pages = len(pdf_reader.pages)

        os.unlink(tmp_docx_path)
        os.unlink(pdf_path)

        return num_pages

    def _update_opis_parameters(self, selected_files):
        for file in selected_files:
            task = file.task_id
            if not task:
                continue

            opis_param = task.opis_parameter_ids.filtered(
                lambda p: p.document_sablon_id.id == file.attachment_id.id
            )
            if opis_param:
                opis_param.write({
                    'denumire_document_opis': file.denumire_document_opis,
                    'prefix_document': file.prefix_document,
                    'cod_document': file.cod_document,
                    'data_document': file.data_document,
                })

    def _clear_existing_attachments(self):
        project = self.project_id
        if not project:
            return

        for task in project.task_ids:
            attachments = self.env['ir.attachment'].search([
                ('res_model', '=', 'project.task'),
                ('res_id', '=', task.id)
            ])
            attachments.unlink()
            for template in task.document_template_ids:
                template.render(task.id, task._name)

    def action_render_selected(self):
        selected_files = self.file_selection_ids.filtered(lambda f: f.selected).sorted("sequence")
        if not selected_files:
            raise ValidationError("Nu există fișiere selectate pentru randare.")
        self._update_opis_parameters(selected_files)
        self._clear_existing_attachments()
        table_of_contents_doc = Document()
        table_of_contents_doc.add_heading("Cuprins", level=1)
        toc_table = table_of_contents_doc.add_table(rows=1, cols=6)
        toc_table.style = 'Table Grid'
        hdr_cells = toc_table.rows[0].cells
        hdr_cells[0].text = "Nr."
        hdr_cells[1].text = "Nume fișier"
        hdr_cells[2].text = "Prefix document"
        hdr_cells[3].text = "Cod document"
        hdr_cells[4].text = "Data document"
        hdr_cells[5].text = "Pagini"

        merged_doc = None
        composer = None
        current_page = 1

        for idx, file in enumerate(selected_files, start=1):
            attachment = file.attachment_id
            if attachment.mimetype != 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                raise ValidationError(f"Fișierul '{attachment.name}' nu este un document Word valid.")

            try:
                document_bytes = base64.b64decode(attachment.datas)
                num_pages = self.get_pdf_page_count_from_docx(document_bytes)
                document = Document(BytesIO(document_bytes))
            except Exception as e:
                raise ValidationError(f"Eroare la procesarea fișierului '{attachment.name}': {e}")

            row_cells = toc_table.add_row().cells
            row_cells[0].text = str(idx)
            row_cells[1].text = file.denumire_document_opis
            row_cells[2].text = file.prefix_document
            row_cells[3].text = file.cod_document
            row_cells[4].text = str(file.data_document)
            row_cells[5].text = f"{current_page} - {current_page + num_pages - 1}"

            current_page += num_pages

            record = self.env[file.task_id._name].browse(file.task_id.id)
            replacements = {f'{{{{{param.key}}}}}': param.value if param.value else '' for param in record.parameter_ids}
            replacements.update({
                '[[prefix_document]]': file.prefix_document or '',
                '[[cod_document]]': file.cod_document or '',
                '[[data_document]]': file.data_document.strftime('%d.%m.%Y') if file.data_document else ''
            })
            for paragraph in document.paragraphs:
                self.replace_placeholder_across_runs(paragraph, replacements)
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            self.replace_placeholder_across_runs(paragraph, replacements)

            if merged_doc is None:
                merged_doc = document
                composer = Composer(merged_doc)
            else:
                composer.doc.add_section(WD_SECTION.NEW_PAGE)
                composer.append(document)
                last_paragraph = composer.doc.paragraphs[-1]

        # Save table of contents
        toc_stream = BytesIO()
        table_of_contents_doc.save(toc_stream)
        toc_stream.seek(0)

        res_model = self.project_id._name
        res_id = self.project_id.id

        self.env['ir.attachment'].create({
            'name': f"{self.project_id.name} - OPIS.docx",
            'type': 'binary',
            'datas': base64.b64encode(toc_stream.read()),
            'res_model': res_model,
            'res_id': res_id,
        })

        # Save merged document
        if merged_doc:
            merged_stream = BytesIO()
            merged_doc.save(merged_stream)
            merged_stream.seek(0)
            self.env['ir.attachment'].create({
                'name': f"{self.project_id.name} - Cartea Constructorului.docx",
                'type': 'binary',
                'datas': base64.b64encode(merged_stream.read()),
                'res_model': res_model,
                'res_id': res_id,
            })

        return True


class DocumentRenderWizardFileSelection(models.TransientModel):
    _name = 'document.render.wizard.file.selection'
    _description = 'File Selection for Render Wizard'

    wizard_id = fields.Many2one('document.render.wizard', string="Wizard", ondelete='cascade')
    name = fields.Char(string="File Name", required=True)
    denumire_document_opis = fields.Char(string="Denumire document Opis")
    prefix_document = fields.Char(string="Prefix document")
    cod_document = fields.Char(string="Cod document")
    data_document = fields.Date(string="Data document")
    attachment_id = fields.Many2one('ir.attachment', string="Attachment")
    selected = fields.Boolean(string="Select", default=False)
    project_id = fields.Many2one("project.project", string="Proiect")
    task_id = fields.Many2one("project.task", string="Sarcina")
    document_template_id = fields.Many2one(
        'document.template',
        string='Document Template',
    )
    sequence = fields.Integer()
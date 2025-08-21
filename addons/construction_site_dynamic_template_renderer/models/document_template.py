from odoo import models, fields, api
from odoo.exceptions import ValidationError
from io import BytesIO
from docx import Document
import zipfile
import base64
import re
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docxcompose.composer import Composer


class DocumentTemplate(models.Model):
    _name = 'document.template'
    _description = 'Document Template'

    name = fields.Char(string="Name", required=True)
    attachment_ids = fields.Many2many(
        'ir.attachment',
        string="Attachments",
        help="Upload files related to this record."
    )
    parameter_ids = fields.One2many(
        comodel_name='document.template.parameter',
        inverse_name='document_template_id',
        string="Parameters"
    )
    prefix_document = fields.Char("Prefix Document")

    @api.model
    def create(self, vals):
        self._validate_attachments(vals.get('attachment_ids', []))
        for record in self:
            if 'attachment_ids' in vals:
                for att in record.attachment_ids:
                    att.document_template_id = record.id
        return super(DocumentTemplate, self).create(vals)

    def write(self, vals):
        if 'attachment_ids' in vals:
            self._validate_attachments(vals['attachment_ids'])
            for record in self:
                for att in record.attachment_ids:
                    att.document_template_id = record.id
        return super(DocumentTemplate, self).write(vals)

    def _validate_attachments(self, attachment_data):
        for attachment in attachment_data:
            attachment = self.env['ir.attachment'].browse(attachment[1])
            if attachment.mimetype != 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                raise ValidationError(f"The file '{attachment.name}' is not a valid")

    def extract_text_from_docx(self, document):
        text = []

        # Extract paragraphs
        for paragraph in document.paragraphs:
            text.append(paragraph.text)

        # Extract text from tables
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    text.append(cell.text)

        return "\n".join(text)

    def extract_variables(self):
        variable_pattern = re.compile(r'\{\{\s*(.*?)\s*\}\}')

        for attachment in self.attachment_ids:
            if attachment.mimetype != 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                continue

            try:
                document_data = base64.b64decode(attachment.datas)
                document = Document(BytesIO(document_data))
            except Exception as e:
                raise ValidationError(f"Nu se poate procesa fișierul '{attachment.name}': {e}")

            found_variables = set()
            all_text = self.extract_text_from_docx(document)
            matches = variable_pattern.findall(all_text)
            found_variables.update(matches)

            for variable in found_variables:
                if not self.parameter_ids.filtered(lambda p: p.key == variable):
                    self.env['document.template.parameter'].create({
                        'document_template_id': self.id,
                        'key': variable,
                        'value': ''
                    })

    def replace_placeholder_across_runs(self, paragraph, replacements):
        runs = paragraph.runs
        if not runs:
            return

        full_text = "".join(run.text for run in runs)

        for key, value in replacements.items():
            start_idx = full_text.find(key)
            while start_idx != -1:
                end_idx = start_idx + len(key)

                char_count = 0
                run_indexes = []
                for i, run in enumerate(runs):
                    run_len = len(run.text)
                    run_start = char_count
                    run_end = char_count + run_len
                    char_count += run_len

                    if run_end > start_idx and run_start < end_idx:
                        run_indexes.append(i)

                if not run_indexes:
                    break

                first_run = runs[run_indexes[0]]
                last_run = runs[run_indexes[-1]]

                involved_text = "".join(runs[i].text for i in run_indexes)
                replaced_text = involved_text.replace(key, value)
                first_run.text = replaced_text

                for i in run_indexes[1:]:
                    runs[i].text = ""

                full_text = full_text[:start_idx] + value + full_text[end_idx:]
                start_idx = full_text.find(key, start_idx + len(value))

    def render(self, record_id, record_model):
        # if not self.attachment_ids:
        #     raise ValueError("Nu există un șablon de document încărcat.")

        record = self.env[record_model].browse(record_id)

        for opis_param in record.opis_parameter_ids:
            try:
                document_bytes = base64.b64decode(opis_param.document_sablon_id.datas)
            except Exception as e:
                raise ValueError(f"Fișierul nu a putut fi decodat: {e}")

            try:
                with zipfile.ZipFile(BytesIO(document_bytes), 'r') as zip_file:
                    if 'word/document.xml' not in zip_file.namelist():
                        raise ValueError("Structura fișierului .docx nu este validă.")
                document = Document(BytesIO(document_bytes))
            except zipfile.BadZipFile:
                raise ValueError("Fișierul din șablon nu este un fișier .docx valid.")
            replacements = {f'{{{{{param.key}}}}}': param.value if param.value else '' for param in record.parameter_ids}
            replacements.update({
                '[[prefix_document]]': opis_param.prefix_document or '',
                '[[cod_document]]': opis_param.cod_document or '',
                '[[data_document]]': opis_param.data_document.strftime('%d.%m.%Y') if opis_param.data_document else ''
            })
            for paragraph in document.paragraphs:
                self.replace_placeholder_across_runs(paragraph, replacements)

            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            self.replace_placeholder_across_runs(paragraph, replacements)
            output_stream = BytesIO()
            document.save(output_stream)
            output_stream.seek(0)

            cod = opis_param.cod_document or ''
            prefix = opis_param.prefix_document or ''
            denumire = opis_param.denumire_document_opis or ''
            file_name = f"{cod}_{prefix}_{denumire}.docx"

            self.env['ir.attachment'].create({
                'name': file_name,
                'type': 'binary',
                'datas': base64.b64encode(output_stream.read()),
                'res_model': record_model,
                'res_id': record_id,
                'attachment_source': 'processed',
                'document_template_id': self.id,
            })
        return True

class DocumentTemplateParameter(models.Model):
    _name = 'document.template.parameter'
    _description = 'Document Template Parameter'

    project_project_id = fields.Many2one(
        comodel_name='project.project',
        string="Project",
    )

    project_task_id = fields.Many2one(
        comodel_name='project.task',
        string="Task",
    )

    document_template_id = fields.Many2one(
        comodel_name='document.template',
        string="Document Template",
        ondelete='cascade'
    )
    key = fields.Char(string="Key", required=True)
    value = fields.Char(string="Value")